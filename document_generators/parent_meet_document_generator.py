import concurrent.futures
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from functions import document_options as doc_opt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_BREAK, WD_PARAGRAPH_ALIGNMENT


def make_parent_meet_documents(mentor_name: str, meeting_date: str, mentee_data: list):
    with concurrent.futures.ThreadPoolExecutor() as executor:
        future1 = executor.submit(mk_parent_visit_pg, mentor_name, meeting_date, mentee_data)
        future2 = executor.submit(mk_parent_feedback_pg, mentor_name, meeting_date, mentee_data)
        doc1 = future1.result()
        doc2 = future2.result()
    merged_doc = Document()
    for element in doc1.element.body:
        merged_doc.element.body.append(element)
    for element in doc2.element.body[1:]:
        merged_doc.element.body.append(element)
    return merged_doc


def mk_parent_visit_pg(mentor_name: str, meeting_date: str, mentee_data: list):
    doc_object = Document()
    doc_opt.make_heading_for_doc(doc_object, 'Parent Visit Form', mentor_name, meeting_date)
    doc_object.add_paragraph('\n')
    doc_opt.make_name_parent_table_for_doc(doc_object, mentee_data)
    purpose_line = doc_object.add_paragraph('\n')
    purpose_line.add_run("Purpose of visit: ").bold = True
    purpose_line.add_run("Parent-Teacher-Meeting/Other (In case of other please mention below)\n\n")
    purpose_line.add_run(2 * "____________________________________________________________________________________\n\n")
    signature_table = doc_object.add_table(rows=2, cols=2)
    signature_table_data = [("____________________________", "____________________________"),
                            (f"{mentor_name}", "Parent/Guardian")]
    for row_index, data_row in enumerate(signature_table_data):
        row = signature_table.rows[row_index]
        for col_index, cell_value in enumerate(data_row):
            cell = row.cells[col_index]
            cell.text = cell_value
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if row_index == 1:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    run = doc_object.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)
    return doc_object


def mk_parent_feedback_pg(mentor_name: str, meeting_date: str, mentee_data: list):
    doc_object = Document()
    doc_object.add_paragraph("")
    doc_opt.make_heading_for_doc(doc_object, 'Parent Feedback Form', mentor_name, meeting_date)
    feedback_request_heading = "Dear parent,\n"
    feedback_heading = doc_object.add_paragraph('\n')
    feedback_heading.add_run(feedback_request_heading)
    feedback_request_content = "We hereby request you to spend your valuable time and provide a rating of " \
                               "the following points in the table provided. Your suggestions and feedback shall go a " \
                               "long way in improving the teaching learning process."
    feedback_request = doc_object.add_paragraph('')
    feedback_request.add_run(feedback_request_content)
    feedback_request.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc_object.add_paragraph('\n')
    feedback_table = doc_object.add_table(rows=13, cols=7, style='Table Grid')
    first_row_content = ('Sr. No', 'Factors for Assessment', 'Excellent', 'Very Good', 'Good', 'Satisfactory', 'Poor')
    factors = ('Teaching Quality', 'General Discipline', 'Library Facility', 'Laboratory Facility',
               'Infrastructural Facility', "Student's Attendance Monitoring", 'Protection Against Ragging',
               'Approachability of Faculty', 'Canteen/Cafeteria Facility', 'Co-Curricular Facilities',
               'Placement Activities', 'Overall Impression')
    for col_index, content in enumerate(first_row_content):
        row = feedback_table.rows[0]
        row.cells[col_index].text = content
    for i in range(1, 13):
        row = feedback_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = factors[i - 1]
    for row in feedback_table.rows:
        for col in range(7):
            cell = row.cells[col]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc_opt.set_table_widths(feedback_table, (0.93, 5.45, 1.9, 1.66, 1.47, 2.34, 1.86))
    doc_object.add_paragraph('\n')
    doc_opt.make_name_parent_table_for_doc(doc_object, mentee_data, guardian_para_flag=False)
    other_sug = doc_object.add_paragraph('\n')
    other_sug.add_run("Any Other Suggestion: ").bold = True
    other_sug.add_run("____________________________________________________________\n\n")
    other_sug.add_run("________________________________________________________________________________________\n")
    doc_opt.add_signature_line(doc_object, '', False, '', False, True)
    doc_opt.line_space_setter(doc_object)
    return doc_object


def super_merger(doc_list: list, filepath: str, meeting_date: str):
    merged_document = Document()
    for doc in doc_list:
        for element in doc.element.body:
            merged_document.element.body.append(element)
    doc_opt.make_header_for_doc(merged_document)
    merged_document.save(f'{filepath}/{meeting_date} Parent Meet Forms.docx')


def ptm_attendance_generator(mentor_name: str, filepath: str, meeting_date: str, mentee_data: list):
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Parent Teacher Meet Attendance', mentor_name, meeting_date)
    names = [item[0] for item in mentee_data]
    roll_nos = [item[1] for item in mentee_data]
    divisions = [item[2] for item in mentee_data]
    doc.add_paragraph("\n")
    ptm_attendance_table = doc.add_table(rows=22, cols=7, style='Table Grid')
    cells_to_merge = (2, 6)
    table_heading_1 = ('Name', 'Signature')
    for i, headings in zip(cells_to_merge, table_heading_1):
        ptm_attendance_table.cell(0, i).merge(ptm_attendance_table.cell(1, i))
        ptm_attendance_table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        header_paragraph = ptm_attendance_table.rows[0].cells[i].paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_paragraph.add_run(headings).bold = True
    detail_headings = ('Division', 'Roll No')
    for i, detail_heading in zip(range(2), detail_headings):
        ptm_attendance_table_content = ptm_attendance_table.rows[1].cells[i].paragraphs[0]
        ptm_attendance_table_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ptm_attendance_table_content.add_run(detail_heading).bold = True
    ptm_attendance_table.cell(0, 0).merge(ptm_attendance_table.cell(0, 1))
    ptm_attendance_table.cell(0, 3).merge(ptm_attendance_table.cell(0, 5))
    table_heading_2 = ('Current Detail', 'Parent/Guardian')
    for i, heading in zip([0, 3], table_heading_2):
        ptm_attendance_table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        ptm_attendance_table_content = ptm_attendance_table.rows[0].cells[i].paragraphs[0]
        ptm_attendance_table_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ptm_attendance_table_content.add_run(heading).bold = True
    detail_parent = ('Mother', 'Father', 'Guardian')
    for i, detail_heading in zip(range(3, 6), detail_parent):
        ptm_attendance_table_content = ptm_attendance_table.rows[1].cells[i].paragraphs[0]
        ptm_attendance_table_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ptm_attendance_table_content.add_run(detail_heading).bold = True
    row_index_range = len(names)+2
    for row_index in range(2, row_index_range):
        name_cell = ptm_attendance_table.cell(row_index, 2)
        name_cell.text = names[row_index - 2]
        division_cell = ptm_attendance_table.cell(row_index, 1)
        division_cell.text = divisions[row_index - 2]
        roll_no_cell = ptm_attendance_table.cell(row_index, 0)
        roll_no_cell.text = roll_nos[row_index - 2]
        for col1_para, col2_para in zip(division_cell.paragraphs, roll_no_cell.paragraphs):
            col1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            col2_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_opt.set_table_widths(ptm_attendance_table, (1.94, 1.75, 5.75, 1.2, 1.2, 1.2, 2.18))
    doc_opt.add_signature_line(doc, mentor_name, True, '', False, False)
    doc.save(f'{filepath}/{meeting_date} PTM Attendance.docx')
