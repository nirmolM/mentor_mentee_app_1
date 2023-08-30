from functions import document_options as doc_opt
from docx import Document
from datetime import date
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def make_academic_achievement_record(academic_achievement_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Academic Achievement Record', mentor_name, today.strftime("%B %d, %Y"))
    academic_achievement_paragraph = doc.add_paragraph("\n")
    academic_achievement_paragraph.add_run("\nAcademic Achievement Record for ")
    academic_achievement_paragraph.add_run(f"{academic_achievement_details['name']}").bold = True
    academic_achievement_paragraph.add_run(": ")
    academic_achievement_table = doc.add_table(rows=6, cols=2)
    academic_achievement_table_headers = ('Achievement Type/Competition Name', 'Subject', 'Semester', 'Year', 'Rank',
                                          'Academic Year')
    academic_achievement_table_details = (academic_achievement_details['achievement_type'],
                                          academic_achievement_details['subject'],
                                          academic_achievement_details['semester'],
                                          academic_achievement_details['year'],
                                          academic_achievement_details['rank'],
                                          academic_achievement_details['academic_year'])
    for i in range(6):
        col = academic_achievement_table.rows[i]
        col.cells[0].text = academic_achievement_table_headers[i]
        col.cells[1].text = academic_achievement_table_details[i]
    doc_opt.set_table_widths(academic_achievement_table, (6, 9.9))
    for row in academic_achievement_table.rows:
        for col in range(2):
            cell = row.cells[col]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc_opt.add_signature_line(doc, mentor_name, True, '', False, False)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{academic_achievement_details['name']} Achievement Record {today.strftime('%B %d, %Y')}.docx")


def make_lor_loa_record(lor_loa_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'LoR/LoA Record', mentor_name, today.strftime("%B %d, %Y"))
    lor_loa_paragraph = doc.add_paragraph("\n")
    lor_loa_paragraph.add_run("\nAcademic Achievement Record for ")
    lor_loa_paragraph.add_run(f"{lor_loa_details['name']}").bold = True
    lor_loa_paragraph.add_run(": ")
    lor_loa_table = doc.add_table(rows=3, cols=2)
    lor_loa_table_headers = ('Letter Type', 'Issuing Faculty', 'Reason')
    lor_loa_table_details = (lor_loa_details['letter_type'],
                             lor_loa_details['issuing_faculty'],
                             lor_loa_details['reason'])
    for i in range(3):
        col = lor_loa_table.rows[i]
        col.cells[0].text = lor_loa_table_headers[i]
        col.cells[1].text = lor_loa_table_details[i]
    doc_opt.set_table_widths(lor_loa_table, (6, 9.9))
    for row in lor_loa_table.rows:
        for col in range(2):
            cell = row.cells[col]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc_opt.add_signature_line(doc, mentor_name, True, '', False, False)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{lor_loa_details['name']} LoR LoA Record {today.strftime('%B %d, %Y')}.docx")
