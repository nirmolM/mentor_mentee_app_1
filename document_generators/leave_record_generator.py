from functions import document_options as doc_opt
from docx import Document
from datetime import date
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def make_leave_record(leave_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Absence Record', mentor_name, today.strftime("%B %d, %Y"))
    leave_paragraph = doc.add_paragraph("\n")
    leave_paragraph.add_run("\nLeave Record for ")
    leave_paragraph.add_run(f"{leave_details['name']}").bold = True
    leave_paragraph.add_run(": ")
    leave_table = doc.add_table(rows=4, cols=2)
    leave_table_headers = ('Starting Date of Absence', 'Absent Till', 'No of Days', 'Reason')
    leave_table_details = (leave_details['start_date'], leave_details['end_date'], str(leave_details['duration']),
                           leave_details['reason'])
    for i in range(4):
        col = leave_table.rows[i]
        col.cells[0].text = leave_table_headers[i]
        col.cells[1].text = leave_table_details[i]
    doc_opt.set_table_widths(leave_table, (4.49, 11.41))
    for row in leave_table.rows:
        for col in range(2):
            cell = row.cells[col]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    description_line = doc.add_paragraph("\n")
    description_line.add_run("Description: ").bold = True
    description_content = doc.add_paragraph("\n")
    description_content.add_run(leave_details['description'])
    description_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc_opt.add_signature_line(doc, mentor_name, True, leave_details['name'], True, False)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{leave_details['name']} Leave Record {today.strftime('%B %d, %Y')}.docx")
