from functions import document_options as doc_opt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def make_attendance_sheet(mentor_name: str, meeting_agenda: str, meeting_date: str, mentee_data: list):
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Mentor-Mentee Meeting Attendance Sheet', mentor_name, meeting_date)
    agenda_line = doc.add_paragraph('')
    agenda_line.add_run(f'\nAgenda: {meeting_agenda}').bold = True
    agenda_line.alignment = 1
    attendance_table = doc.add_table(rows=int(len(mentee_data)/3)+2, cols=5, style='Table Grid')
    attendance_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_nos_for_merging = (0, 1, 4)
    table_headings = ('Sr. No', 'Name', 'Signature')
    for i, headings in zip(cell_nos_for_merging, table_headings):
        attendance_table.cell(0, i).merge(attendance_table.cell(1, i))
        attendance_table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        header_paragraph = attendance_table.rows[0].cells[i].paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_paragraph.add_run(headings).bold = True
    attendance_table.cell(0, 2).merge(attendance_table.cell(0, 3))
    attendance_table.cell(0, 2).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    attendance_table_content = attendance_table.rows[0].cells[2].paragraphs[0]
    attendance_table_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    attendance_table_content.add_run('Current Detail').bold = True
    detail_headings = ('Division', 'Roll No')
    for i, detail_heading in zip(range(2, 4), detail_headings):
        attendance_table_content = attendance_table.rows[1].cells[i].paragraphs[0]
        attendance_table_content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        attendance_table_content.add_run(detail_heading).bold = True
    doc_opt.set_table_widths(attendance_table, (0.87, 8.83, 3.26, 2.67))
    sr_no = list(range(1, int(len(mentee_data)/3)+1))
    names = mentee_data[0::3]
    divisions = mentee_data[1::3]
    roll_nos = mentee_data[2::3]
    row_index_range = list(range(2, int(len(mentee_data)/3)+2))
    for row_index in row_index_range:
        sr_no_cell = attendance_table.cell(row_index, 0)
        sr_no_cell.text = str(sr_no[row_index - 2])
        name_cell = attendance_table.cell(row_index, 1)
        name_cell.text = names[row_index - 2]
        division_cell = attendance_table.cell(row_index, 2)
        division_cell.text = divisions[row_index - 2]
        roll_no_cell = attendance_table.cell(row_index, 3)
        roll_no_cell.text = roll_nos[row_index - 2]
        for col1_para, col2_para, col3_para in zip(sr_no_cell.paragraphs, division_cell.paragraphs,
                                                   roll_no_cell.paragraphs):
            col1_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            col2_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            col3_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_opt.add_signature_line(doc, mentor_name, True, '', False, False)
    return doc
