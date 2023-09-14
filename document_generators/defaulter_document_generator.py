from functions import document_options as doc_opt
from working_data import year_semester_giver as ysg
from docx import Document
from datetime import date
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def parent_letter_generator(defaulter_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    to_para = doc.add_paragraph()
    to_para.add_run("To,").bold = True
    father_name_para = doc.add_paragraph("\n")
    father_name_para.add_run(defaulter_details['father_name']).bold = True
    address_para = doc.add_paragraph("\n")
    address_para.add_run(f"{defaulter_details['address']}\n")
    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc_opt.make_heading_for_doc(doc, 'Sub: Attendance of your ward in theory classes & Practical sessions',
                                 mentor_name, today.strftime("%B %d, %Y"))
    dear_parent_line = doc.add_paragraph("\n\n")
    dear_parent_line.add_run('Dear Parent,')
    content_para1 = doc.add_paragraph("\n")
    content_para1.add_run(f"This is to inform you that your ward {defaulter_details['name']}, "
                          f"is having less attendance "
                          f"({defaulter_details['attendance']}) in theory classes, practical sessions and tutorials.")
    content_para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    content_para2 = doc.add_paragraph("\n")
    content_para2.add_run("As per the rule of Mumbai University, your ward will not be allowed to appear for the "
                          "examination if he/she fails to maintain 75% attendance in the semester.")
    content_para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    content_para3 = doc.add_paragraph()
    content_para3.add_run("We request you to immediately look into your ward’s attendance with the highest priority.")
    content_para3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    hod_signature_line = doc.add_paragraph("\n\n")
    hod_signature_line.add_run("_________________________________________\n")
    hod_signature_line.add_run("Head of Department").bold = True
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{defaulter_details['name']} Parent Letter {today.strftime('%B %d, %Y')}.docx")


def undertaking_student(defaulter_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, '\nUNDERTAKING', mentor_name, today.strftime("%B %d, %Y"))
    subject_line = doc.add_paragraph("\n")
    subject_line.add_run("Subject: Attendance of student in theory classes and practical sessions.").bold = True
    subject_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_line = doc.add_paragraph("\n")
    name_line.add_run(f"Name: {defaulter_details['name']}").bold = True
    division_line = doc.add_paragraph()
    division_line.add_run(f"Class/Division: {defaulter_details['division']}").bold = True
    roll_no_line = doc.add_paragraph()
    roll_no_line.add_run(f"Roll No: {defaulter_details['roll_no']}").bold = True
    semester_line = doc.add_paragraph()
    semester_line.add_run(f"Semester: {ysg.give_year_and_semester(defaulter_details['admitted_year'])[1]}").bold = True
    content_para1 = doc.add_paragraph("\n")
    content_para1.add_run(f"I, {defaulter_details['name']}, the undersigned, student of this college, am fully aware "
                          f"that I am required to attend classes on all the working days and due to any genuine "
                          f"reason, if I am not in a position to attend on any day, I shall inform my HOD in writing "
                          f"in advance and take her permission to remain absent. If I could not attend on any day, "
                          f"I shall inform on next working day.")
    content_para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    content_para2 = doc.add_paragraph("\n")
    content_para2.add_run("I also clearly understand that I will not be allowed to appear for the examination if I "
                          "don’t attend at least minimum 75% classes of theory, tutorials and practical including "
                          "drawing, I am also aware that I will not be allowed to appear for examination, if I fail to "
                          "submit within stipulated time limit all the assignments, jobs, journals, drawing etc. "
                          "satisfactorily, as specific by the university.")
    content_para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc_opt.add_signature_line(doc, '', False, defaulter_details['name'], True, True)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{defaulter_details['name']} Undertaking {today.strftime('%B %d, %Y')}.docx")
