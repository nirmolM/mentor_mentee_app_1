from functions import document_options as doc_opt
from docx import Document
from datetime import date
from working_data import year_semester_giver as ysg
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def make_special_action_sheet(issue_details: dict, filepath: str, mentor_name: str):
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Mentor-Mentee Issue', mentor_name,
                                 date.today().strftime("%B %d, %Y"))
    doc.add_paragraph('')
    issue_table = doc.add_table(rows=10, cols=4, style='Table Grid')
    issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_to_be_merged = (0, 3, 4, 5, 6, 7, 8, 9)
    for i in rows_to_be_merged:
        issue_table.cell(i, 0).merge(issue_table.cell(i, 3))
    academic_year_line = issue_table.rows[0].cells[0].paragraphs[0]
    academic_year_line.add_run("Academic Year: ").bold = True
    academic_year_line.add_run(ysg.give_academic_year())
    issue_table.cell(1, 0).merge(issue_table.cell(2, 0))
    mentee_details_paragraph = issue_table.rows[1].cells[0].paragraphs[0]
    mentee_details_paragraph.add_run("Mentee Details: ").bold = True
    mentee_details_headings = ('Name', 'Division', 'Roll No')
    mentee_details_content = (issue_details['name'], issue_details['division'], issue_details['roll_no'])
    cells_in_zip = (1, 2, 3)
    for i, headings, content in zip(cells_in_zip, mentee_details_headings, mentee_details_content):
        header_paragraph = issue_table.rows[1].cells[i].paragraphs[0]
        header_paragraph.add_run(headings).bold = True
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        content_paragraph = issue_table.rows[2].cells[i].paragraphs[0]
        content_paragraph.add_run(content)
        content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    issue_raised_line = issue_table.rows[3].cells[0].paragraphs[0]
    issue_raised_line.add_run("Issue Raised: ").bold = True
    issue_raised_line.add_run(issue_details['issue'])
    issue_raised_line.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    issue_raised_content = issue_table.rows[4].cells[0].paragraphs[0]
    issue_raised_content.add_run(issue_details['issue_description'])
    issue_raised_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    issue_resolved_line = issue_table.rows[5].cells[0].paragraphs[0]
    issue_resolved_line.add_run("Issue Resolved: ").bold = True
    action_taken_line = issue_table.rows[6].cells[0].paragraphs[0]
    action_taken_line.add_run("Action Taken: ").bold = True
    action_taken_input = issue_table.rows[7].cells[0].paragraphs[0]
    action_taken_input.add_run(issue_details['action'])
    action_taken_input.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    outcome_line = issue_table.rows[8].cells[0].paragraphs[0]
    outcome_line.add_run("Improvement/Outcome: ").bold = True
    outcome_line_input = issue_table.rows[9].cells[0].paragraphs[0]
    outcome_line_input.add_run(issue_details['outcome'])
    outcome_line_input.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for i in range(10):
        issue_table.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc_opt.add_signature_line(doc, mentor_name, True, issue_details['name'], True, False)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{issue_details['name']} Mentor-Mentee Action {date.today().strftime('%B %d, %Y')}.docx")
