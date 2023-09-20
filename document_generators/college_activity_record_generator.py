from functions import document_options as doc_opt
from docx import Document
from datetime import date
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def make_college_activity_record(activity_details: dict, filepath: str, mentor_name: str):
    today = date.today()
    doc = Document()
    doc_opt.make_header_for_doc(doc)
    doc_opt.make_heading_for_doc(doc, 'Activity Record', mentor_name, today.strftime("%B %d, %Y"))
    activity_paragraph = doc.add_paragraph("\n")
    activity_paragraph.add_run("\nActivity Record for ")
    activity_paragraph.add_run(f"{activity_details['name']}").bold = True
    activity_paragraph.add_run(": ")
    activity_table = doc.add_table(rows=4, cols=2)
    activity_table_headers = ('Activity Type', 'Role', 'Day of Event', 'End Day of Event')
    activity_table_details = (activity_details['activity_type'], activity_details['role'],
                              activity_details['start_date'], activity_details['end_date'])
    for i in range(4):
        col = activity_table.rows[i]
        col.cells[0].text = activity_table_headers[i]
        col.cells[1].text = activity_table_details[i]
    doc_opt.set_table_widths(activity_table, (4.49, 11.41))
    for row in activity_table.rows:
        for col in range(2):
            cell = row.cells[col]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    description_line = doc.add_paragraph("\n")
    description_line.add_run("Description: ").bold = True
    description_content = doc.add_paragraph("\n")
    description_content.add_run(activity_details['description'])
    description_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc_opt.add_signature_line(doc, mentor_name, True, '', False, False)
    doc_opt.line_space_setter(doc)
    doc.save(f"{filepath}/{activity_details['name']} Leave Record {today.strftime('%B %d, %Y')}.docx")
