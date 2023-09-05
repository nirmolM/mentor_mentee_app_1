from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def make_header_for_doc(doc_object):
    header_object = doc_object.sections[0].header
    header_table = header_object.add_table(1, 2, Inches(8))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2):
        header_table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        header_paragraph = header_table.rows[0].cells[i].paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_paragraph = header_table.rows[0].cells[0].paragraphs[0]
    header_content = header_paragraph.add_run()
    header_content.add_picture('document_generators/Header_image.png', width=Inches(1), height=Inches(1))
    header_table.columns[0].width = Inches(1.3)
    header_table.rows[0].cells[0].width = Inches(1.3)
    header_line1 = "Mahavir Education's Trust\n"
    header_line2 = "SHAH & ANCHOR KUTCHHI ENGINEERING COLLEGE\n"
    header_line3 = "Chembur, 400088\n"
    header_line4 = "U.G. Program in Electronics and Computer Science"
    header_lines = header_table.rows[0].cells[1].paragraphs[0]
    header_lines.add_run(header_line1 + header_line2 + header_line3 + header_line4)
    header_row_cells = header_table.rows[0].cells
    for cell in header_row_cells:
        cell.alignment = 1
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def make_heading_for_doc(doc_object, heading_string: str, mentor_name: str, meeting_date: str):
    heading = doc_object.add_heading(heading_string, level=1)
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run.bold = True
    heading_run.font.name = 'Times New Roman'
    meeting_info_table = doc_object.add_table(1, 2)
    meeting_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_cells = meeting_info_table.rows[0].cells
    strings = (mentor_name, meeting_date)
    for cell, i in zip(row_cells, range(len(strings))):
        cell.alignment = 1
        cell.text = strings[i]
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def make_name_parent_table_for_doc(doc_object, mentee_data, guardian_para_flag=True):
    name_table = doc_object.add_table(rows=3, cols=2)
    name_table_col1 = ('Name: ', 'Roll No: ', 'Class: ')
    name_table_col2 = (mentee_data[0], mentee_data[1], mentee_data[2])
    name_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        row = name_table.rows[i]
        row.cells[0].text = name_table_col1[i]
        row.cells[1].text = name_table_col2[i]
        run = row.cells[0].paragraphs[0].runs[0]
        run.bold = True
        row_cells = name_table.rows[i].cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    widths_1 = (2, 13.24)
    for i, width in enumerate(widths_1):
        name_table.columns[i].width = Cm(width)
        name_table.rows[i].cells[i].width = Cm(width)
    parent_table = doc_object.add_table(rows=4, cols=4)
    parent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data_to_add = [("", "Mother", "Father", "Guardian"),
                   ("Name: ", mentee_data[3], mentee_data[6], mentee_data[9]),
                   ("Contact No: ", mentee_data[4], mentee_data[7], mentee_data[10]),
                   ("Email Id: ", mentee_data[5], mentee_data[8], mentee_data[11])]
    for row_index, data_row in enumerate(data_to_add):
        row = parent_table.rows[row_index]
        for col_index, cell_value in enumerate(data_row):
            cell = row.cells[col_index]
            cell.text = cell_value
            if col_index in [1, 2, 3]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if row_index == 0 or col_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    if guardian_para_flag:
        lines = [("Name: ", "_____________________________________"),
                 ("Contact No: ", "_____________________________________"),
                 ("Relation: ", "_____________________________________")]
        doc_object.add_paragraph("\nIf Guardian is not same as above, Please enter details:\n")
        for header, placeholder in lines:
            p = doc_object.add_paragraph()
            p.add_run(header).bold = True
            p.add_run(placeholder + "\n")


def add_signature_line(doc_object,
                       mentor_name: str, mentor_flg: bool, mentee_name: str, mentee_flg: bool, parent_flg: bool):
    if mentor_flg:
        mentor_signature_line = doc_object.add_paragraph("")
        mentor_signature_line.add_run('\n\n______________________')
        mentor_signature_line.add_run(f'\n{mentor_name}').bold = True
    if mentee_flg:
        mentee_signature_line = doc_object.add_paragraph("")
        mentee_signature_line.add_run('\n\n______________________')
        mentee_signature_line.add_run(f'\n{mentee_name}').bold = True
    if parent_flg:
        parent_signature_line = doc_object.add_paragraph()
        parent_signature_line.add_run("\n_________________________________________\n")
        parent_signature_line.add_run("Parent/Guardian").bold = True


def set_table_widths(doc_table_object, widths: tuple):
    for i, width in enumerate(widths):
        doc_table_object.columns[i].width = Cm(width)
        doc_table_object.rows[i].cells[i].width = Cm(width)


def line_space_setter(doc_object):
    for paragraph in doc_object.paragraphs[0:]:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
